"""The dedicated interpretation summary — assembled, reviewable, downloadable.

The Summary page keeps its marker table, pathways result, and cross-cluster
review, and gains this fourth region: a per-dataset interpretation report the
biologist reviews on-page and downloads as .docx and PDF.

It is assembled ENTIRELY from durable, already-grounded artifacts — the per-cluster
verdicts (jazzPanda numbers), the cited cell-type biology notes, the deterministic
"what would settle it" line for shaky calls, the lab's own notes with their
agree/dissent citations, and the cross-cluster holistic review. There is **no live
network at build/download time**: nothing is fetched, nothing is recomputed, so the
report is reproducible and unit-testable, yet it still reads as "what came out of
your conversations with the agent."

Design boundaries:
- ``build_report`` is a pure function of its inputs (no I/O) → fully testable.
- Streamlit / ``ui.data_access`` are imported lazily, so ``build_report`` /
  ``report_to_docx`` / ``report_to_pdf`` work with no server running.
- No new ``ui.data_access`` accessor is added (the concurrent enrichment session
  owns that file's tail); this module calls only committed accessors.
"""
from __future__ import annotations

import html as _html
import re as _re
from dataclasses import dataclass

from agent import config as cfg
from agent import discriminate
from agent.types import ClusterVerdict, MarkerEvidence, Note

# --------------------------------------------------------------------------- #
# Human labels for the closed note vocabularies (relabel only; never change one).
# Inlined here (rather than importing ui.lab_knowledge) so this module's pure
# functions stay free of Streamlit / data_access.
# --------------------------------------------------------------------------- #
_SCOPE_LABEL = {"cluster": "cluster", "dataset": "this dataset", "lab": "lab-wide"}
_BASIS_LABEL = {"paper": "a paper", "own_validation": "our own data", "convention": "convention"}
_STATUS_LABEL = {"firm": "firm rule", "tentative": "tentative"}


# --------------------------------------------------------------------------- #
# Report model — plain, serializable, no I/O
# --------------------------------------------------------------------------- #
@dataclass(frozen=True)
class NoteLine:
    claim: str
    scope: str
    basis: str
    status: str
    tension: str  # "" when none
    author: str
    date: str


@dataclass(frozen=True)
class ClusterSection:
    cluster: str
    cell_type: str
    confidence: str
    verify: bool
    driver_line: str  # "ERBB2 (glm_coef 21.44, pearson 0.59), ..." or key markers
    rationale: str
    biology: str
    biology_pmid: str  # "" when none
    settle: str        # "" unless there is a genuine rival to settle
    notes: tuple[NoteLine, ...]
    enrichment: str = ""             # plain-text enriched-programs block (with PMIDs), "" if none
    enrichment_confidence: str = ""  # the cluster's enrichment confidence band, "" if none


@dataclass(frozen=True)
class ReportModel:
    dataset: str
    generated_at: str  # passed in; "" for a deterministic (test) build
    n_clusters: int
    n_flagged: int
    panel_size: int
    sections: tuple[ClusterSection, ...]
    coherence_notes: tuple[str, ...]
    refinements: tuple[str, ...]
    set_is_coherent: bool
    dataset_notes: tuple[NoteLine, ...]  # dataset/lab-scoped notes, shown once


# --------------------------------------------------------------------------- #
# Assembly (pure)
# --------------------------------------------------------------------------- #
def _pmids(cites) -> str:
    return ", ".join(f"PMID:{c.pmid}" for c in cites if getattr(c, "pmid", ""))


def _tension_text(note: Note) -> str:
    """Literature-tension summary off the stored note, or '' when there is none."""
    t = note.tension
    if t.agree or t.dissent:
        bits = []
        if t.agree:
            bits.append(f"agrees ({len(t.agree)}): {_pmids(t.agree)}")
        if t.dissent:
            bits.append(f"dissents ({len(t.dissent)}): {_pmids(t.dissent)}")
        return "literature tension: " + " · ".join(bits)
    if t.thin:
        return "literature thin: no supporting reference on file"
    return ""


def _note_line(note: Note) -> NoteLine:
    if note.scope == "cluster" and note.scope_ref.cluster:
        scope = f"cluster {note.scope_ref.cluster}"
    else:
        scope = _SCOPE_LABEL.get(note.scope, note.scope)
    return NoteLine(
        claim=note.claim,
        scope=scope,
        basis=_BASIS_LABEL.get(note.basis, note.basis),
        status=_STATUS_LABEL.get(note.status, note.status),
        tension=_tension_text(note),
        author=note.author or "you",
        date=(note.created_at.split("T", 1)[0] if note.created_at else "n/a"),
    )


def _driver_line(v: ClusterVerdict) -> str:
    """Driving markers with real numbers, or the key-marker names if none drive."""
    drivers: tuple[MarkerEvidence, ...] = v.opening.driving_markers
    if drivers:
        return ", ".join(
            f"{m.gene} (glm_coef {m.glm_coef:.2f}, pearson {m.pearson:.2f})" for m in drivers[:4]
        )
    return ", ".join(v.key_markers) if v.key_markers else "no canonical driver"


def _settle_for(cluster: str, verify: bool) -> str:
    """The 'what would settle it' summary, only when there is a genuine rival."""
    d = discriminate.discriminate(cluster)
    if d.alt_B is not None and (verify or d.b_here):
        return discriminate.settle_summary(d)
    return ""


def _enrichment_block(ce, pathway_notes_for_cluster: dict) -> tuple[str, str]:
    """A plain-text block of the cluster's enriched programs + cited biology, and the
    enrichment confidence band. ('' , '') when there is no enrichment for the cluster."""
    if ce is None:
        return "", ""
    if not getattr(ce, "enriched", ()):  # no program clears the gate
        return "", ce.confidence
    bits: list[str] = []
    for p in ce.enriched[:4]:
        short = p.gene_set.replace("HALLMARK_", "").replace("_", " ").title()
        note = (pathway_notes_for_cluster or {}).get(p.gene_set) or {}
        bio = str(note.get("summary") or "").strip()
        pmid = note.get("pmid")
        cite = f" (PMID:{pmid})" if pmid else ""
        bits.append(f"{short} — {bio}{cite}" if bio else short)
    return "; ".join(bits) + ".", ce.confidence


def build_report(
    *,
    verdicts: list[ClusterVerdict],
    celltype_notes: dict,
    notes: list[Note],
    holistic,
    panel_size: int,
    dataset: str = cfg.DATASET_ID,
    generated_at: str = "",
    enrichments: dict | None = None,
    pathway_notes: dict | None = None,
) -> ReportModel:
    """Assemble the interpretation report from durable, grounded inputs. Pure.

    ``enrichments`` maps cluster -> ClusterEnrichment and ``pathway_notes`` maps
    cluster -> {gene_set: note}; when given, each section integrates the cluster's
    enriched programs + their cited biology alongside the marker call.
    """
    enrichments = enrichments or {}
    pathway_notes = pathway_notes or {}
    sections: list[ClusterSection] = []
    for v in verdicts:
        ct_note = celltype_notes.get(v.cluster) or {}
        cluster_notes = tuple(
            _note_line(n)
            for n in notes
            if n.scope == "cluster" and n.scope_ref.cluster == v.cluster
        )
        enr_text, enr_conf = _enrichment_block(
            enrichments.get(v.cluster), pathway_notes.get(v.cluster)
        )
        sections.append(
            ClusterSection(
                cluster=v.cluster,
                cell_type=v.cell_type,
                confidence=v.confidence,
                verify=v.verify,
                driver_line=_driver_line(v),
                rationale=v.notes,
                biology=str(ct_note.get("summary") or ""),
                biology_pmid=str(ct_note.get("pmid") or ""),
                settle=_settle_for(v.cluster, v.verify),
                notes=cluster_notes,
                enrichment=enr_text,
                enrichment_confidence=enr_conf,
            )
        )

    dataset_notes = tuple(_note_line(n) for n in notes if n.scope in ("dataset", "lab"))

    coherence_notes: tuple[str, ...] = ()
    refinements: tuple[str, ...] = ()
    set_is_coherent = True
    if holistic is not None:
        coherence_notes = tuple(holistic.coherence_notes)
        refinements = tuple(
            f"{r.cluster}: {r.from_call} → {r.to_call} — {r.rationale}"
            for r in holistic.refinements
        )
        set_is_coherent = bool(holistic.set_is_coherent)

    return ReportModel(
        dataset=dataset,
        generated_at=generated_at,
        n_clusters=len(sections),
        n_flagged=sum(1 for s in sections if s.verify),
        panel_size=panel_size,
        sections=tuple(sections),
        coherence_notes=coherence_notes,
        refinements=refinements,
        set_is_coherent=set_is_coherent,
        dataset_notes=dataset_notes,
    )


def build_report_from_sources(generated_at: str = "") -> ReportModel:
    """Convenience: pull the durable artifacts from ``ui.data_access`` and assemble.

    Used by the Summary page. Lazily imports data_access so the pure functions above
    (and the doc exporters) never require a Streamlit runtime.
    """
    from ui import data_access as da

    try:
        enrichments = {ce.cluster: ce for ce in da.all_enrichments()}
        pathway_notes = da.pathway_notes()
    except Exception:  # noqa: BLE001 - no enrichment slice -> marker-only report
        enrichments, pathway_notes = {}, {}

    return build_report(
        verdicts=da.composed_verdicts(),  # reflect confirmed overrides/excludes in the report
        celltype_notes=da.celltype_notes(),
        notes=da.read_notes(),
        holistic=da.holistic(),
        panel_size=len(da.panel_names()),
        dataset=cfg.DATASET_ID,
        generated_at=generated_at,
        enrichments=enrichments,
        pathway_notes=pathway_notes,
    )


# --------------------------------------------------------------------------- #
# On-page review region (HTML) — pure builder + Streamlit renderer
# --------------------------------------------------------------------------- #
_PMID_RE = _re.compile(r"PMID:(\d+)")
_CONF_COLOR = {
    "Very High": "#1a7f43",
    "High": "#2f9e57",
    "Medium-High": "#b8860b",
    "Medium": "#c26a1a",
    "Low": "#b23b3b",
}


def _linkify(text: str) -> str:
    """Escape text, then turn PMID:xxxx into a real PubMed link."""
    esc = _html.escape(text)
    return _PMID_RE.sub(
        r'<a href="https://pubmed.ncbi.nlm.nih.gov/\1/" target="_blank">PMID:\1</a>', esc
    )


def _chip(conf: str) -> str:
    color = _CONF_COLOR.get(conf, "#666")
    return (
        f'<span style="background:{color};color:#fff;padding:1px 8px;border-radius:9px;'
        f'font-size:11px;font-weight:600">{_html.escape(conf)}</span>'
    )


def _note_html(nl: NoteLine) -> str:
    tension = (
        f'<div style="font-size:11px;color:var(--muted,#777)">{_linkify(nl.tension)}</div>'
        if nl.tension
        else ""
    )
    return (
        f'<div style="border-left:2px solid var(--accent,#3b6fd4);padding:4px 10px;margin:6px 0">'
        f'<div>{_html.escape(nl.claim)}</div>'
        f'<div style="font-size:11px;color:var(--muted,#777)">'
        f'{_html.escape(nl.scope)} · basis: {_html.escape(nl.basis)} · {_html.escape(nl.status)} · '
        f'{_html.escape(nl.author)} · {_html.escape(nl.date)}</div>{tension}</div>'
    )


def report_html(report: ReportModel) -> str:
    """The interpretation summary as a self-contained HTML string (pure)."""
    parts: list[str] = []
    meta = (
        f"{report.n_clusters} clusters · {report.n_flagged} flagged for re-check · "
        f"{report.panel_size}-gene panel"
    )
    if report.generated_at:
        meta += f" · {_html.escape(report.generated_at)}"
    parts.append(
        f'<div style="font-size:13px"><p style="font-size:12px;color:var(--muted,#777);'
        f'letter-spacing:.04em;text-transform:uppercase">Interpretation summary</p>'
        f'<p style="color:var(--muted,#777);margin:-4px 0 12px">{_html.escape(meta)}</p>'
    )
    for s in report.sections:
        flag = ' <span style="color:#b23b3b;font-size:11px">⚑ re-check</span>' if s.verify else ""
        parts.append(
            f'<div style="border:1px solid var(--hair,#e3e3e3);border-radius:9px;padding:10px 12px;margin:8px 0">'
            f'<div style="display:flex;gap:8px;align-items:center">'
            f'<strong>{_html.escape(s.cluster)} · {_html.escape(s.cell_type)}</strong>{_chip(s.confidence)}{flag}</div>'
            f'<div style="font-size:12px;margin-top:5px"><em>Drivers:</em> {_linkify(s.driver_line)}</div>'
        )
        if s.biology:
            bio = s.biology + (f" (PMID:{s.biology_pmid})" if s.biology_pmid else "")
            parts.append(f'<div style="font-size:12px;margin-top:4px"><em>Biology:</em> {_linkify(bio)}</div>')
        if s.settle:
            parts.append(
                f'<div style="font-size:12px;margin-top:4px;color:var(--accent,#3b6fd4)">'
                f'<em>What would settle it:</em> {_linkify(s.settle)}</div>'
            )
        for nl in s.notes:
            parts.append(_note_html(nl))
        parts.append("</div>")

    # Cross-cluster review recap (the report's own copy of Step 4).
    parts.append('<div style="margin-top:14px"><strong>Cross-cluster review</strong>')
    for c in report.coherence_notes:
        parts.append(f'<div style="font-size:12px;margin-top:4px">{_linkify(c)}</div>')
    for r in report.refinements:
        parts.append(f'<div style="font-size:12px;margin-top:4px">↳ {_linkify(r)}</div>')
    parts.append("</div>")

    if report.dataset_notes:
        parts.append('<div style="margin-top:14px"><strong>Lab notes (dataset / lab-wide)</strong>')
        for nl in report.dataset_notes:
            parts.append(_note_html(nl))
        parts.append("</div>")

    parts.append("</div>")
    return "".join(parts)


def render_report_html(report: ReportModel) -> None:
    """Render the review region into the Summary page."""
    import streamlit as st

    st.markdown(report_html(report), unsafe_allow_html=True)


# --------------------------------------------------------------------------- #
# Downloadable documents (.docx + PDF) — deterministic, no network
# --------------------------------------------------------------------------- #
_PDF_MAP = {
    "—": "-", "–": "-", "→": "->", "←": "<-", "↳": "->", "≈": "~", "×": "x",
    "…": "...", "’": "'", "‘": "'", "“": '"', "”": '"', "•": "*", "·": "-", "⚑": "!",
}


def _pdf_safe(s: str) -> str:
    """fpdf2 core fonts are latin-1; map the unicode punctuation we emit, then
    replace anything still outside latin-1 so a stray glyph can never crash export."""
    for k, v in _PDF_MAP.items():
        s = s.replace(k, v)
    return s.encode("latin-1", "replace").decode("latin-1")


def _note_text(nl: NoteLine) -> str:
    base = f'"{nl.claim}" — {nl.scope}, basis: {nl.basis}, {nl.status} ({nl.author}, {nl.date})'
    return base + (f" — {nl.tension}" if nl.tension else "")


def _section_lines(report: ReportModel) -> list[tuple[str, str]]:
    """Flatten the report into (style, text) lines shared by both exporters.

    style in {title, meta, h2, body, note}. Keeps docx and pdf in sync.
    """
    lines: list[tuple[str, str]] = [("title", "Panoscope — interpretation summary")]
    meta = (
        f"{report.dataset} · {report.n_clusters} clusters · {report.n_flagged} flagged · "
        f"{report.panel_size}-gene panel"
    )
    if report.generated_at:
        meta += f" · {report.generated_at}"
    lines.append(("meta", meta))

    for s in report.sections:
        head = f"{s.cluster} · {s.cell_type} — {s.confidence}" + ("  [re-check]" if s.verify else "")
        lines.append(("h2", head))
        lines.append(("body", f"Drivers: {s.driver_line}"))
        if s.biology:
            bio = s.biology + (f" (PMID:{s.biology_pmid})" if s.biology_pmid else "")
            lines.append(("body", f"Biology: {bio}"))
        if s.settle:
            lines.append(("body", f"What would settle it: {s.settle}"))
        for nl in s.notes:
            lines.append(("note", _note_text(nl)))

    lines.append(("h2", "Cross-cluster review"))
    for c in report.coherence_notes:
        lines.append(("body", c))
    for r in report.refinements:
        lines.append(("body", f"- {r}"))

    if report.dataset_notes:
        lines.append(("h2", "Lab notes (dataset / lab-wide)"))
        for nl in report.dataset_notes:
            lines.append(("note", _note_text(nl)))
    return lines


def _lines_to_docx(lines: list[tuple[str, str]]) -> bytes:
    """Render (style, text) lines as a Word .docx (editable). Requires python-docx."""
    from io import BytesIO

    from docx import Document

    doc = Document()
    for style, text in lines:
        if style == "title":
            doc.add_heading(text, level=0)
        elif style == "meta":
            doc.add_paragraph(text)
        elif style == "h2":
            doc.add_heading(text, level=2)
        elif style == "note":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _lines_to_pdf(lines: list[tuple[str, str]]) -> bytes:
    """Render (style, text) lines as a PDF (locked). Requires fpdf2."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def cell(h: float, text: str) -> None:
        # new_x=LMARGIN resets the cursor to the left margin so the next full-width
        # multi_cell has room (fpdf2 otherwise leaves x at the right margin).
        pdf.multi_cell(0, h, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    for style, raw in lines:
        text = _pdf_safe(raw)
        if style == "title":
            pdf.set_font("Helvetica", "B", 16)
            cell(8, text)
            pdf.ln(1)
        elif style == "meta":
            pdf.set_font("Helvetica", "I", 9)
            cell(5, text)
            pdf.ln(2)
        elif style == "h2":
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 12)
            cell(6, text)
        elif style == "note":
            pdf.set_font("Helvetica", "I", 10)
            cell(5, f"  {text}")
        else:
            pdf.set_font("Helvetica", "", 10)
            cell(5, text)
    return bytes(pdf.output())


def report_to_docx(report: ReportModel) -> bytes:
    """Render the structured report as a Word .docx (editable). Requires python-docx."""
    return _lines_to_docx(_section_lines(report))


def report_to_pdf(report: ReportModel) -> bytes:
    """Render the structured report as a PDF (locked). Requires fpdf2."""
    return _lines_to_pdf(_section_lines(report))


# --------------------------------------------------------------------------- #
# Editable working space — the auto-seed per-cluster summary + its export.
# --------------------------------------------------------------------------- #
def default_cluster_summary(s: ClusterSection) -> str:
    """The auto-seed integrated per-cluster summary (the editable working space).

    Reads like a biologist's synthesis, in two grounded movements: IDENTITY (the
    marker cell-type call + its drivers + cited biology — *what the cell is*) and
    PROGRAMS (the enriched gene sets + their cited biology — *what the cell is
    doing*, panel-scoped), then any what-would-settle-it line and this cluster's
    lab notes. Plain text with PMID:xxx inline (linkified on-page, kept verbatim in
    the exported document). The concordance read (do the programs fit the call?) is
    carried by the per-program biology, which flags cross-lineage signal explicitly.
    """
    ident = (
        f"Identity — {s.cluster} is {s.cell_type} at {s.confidence} confidence"
        + (" (flagged to re-check)" if s.verify else "")
        + f", driven by {s.driver_line}."
    )
    if s.biology:
        ident += " " + s.biology + (f" (PMID:{s.biology_pmid})" if s.biology_pmid else "")
    paras = [ident]

    if s.enrichment:
        conf = f" ({s.enrichment_confidence} enrichment)" if s.enrichment_confidence else ""
        paras.append(
            f"Programs — enriched gene sets{conf}, panel-scoped (measured only over the "
            f"set genes on the panel, not genome-wide): {s.enrichment}"
        )
    elif s.enrichment_confidence:
        paras.append("Programs — no gene-set program clears the enrichment gate for this cluster.")

    if s.settle:
        paras.append(f"What would settle it — {s.settle}")
    for nl in s.notes:
        t = f" — {nl.tension}" if nl.tension else ""
        paras.append(f'Lab note — "{nl.claim}" ({nl.scope} · {nl.basis} · {nl.status}){t}.')
    return "\n\n".join(paras)


def _blocks(text: str) -> list[tuple[str, str]]:
    return [("body", p.strip()) for p in text.split("\n\n") if p.strip()]


def _working_lines(
    sections: list[tuple[str, str, str, bool, str]],
    *,
    dataset: str,
    generated_at: str,
    global_check: str = "",
    caveats: str = "",
    lab_notes: tuple[NoteLine, ...] = (),
) -> list[tuple[str, str]]:
    """(style, text) lines for the EDITED working-space export — per-cluster summaries,
    the cross-cluster global check, then the caveats (the biologist's report shape)."""
    lines: list[tuple[str, str]] = [("title", "Panoscope — interpretation summary")]
    meta = f"{dataset} · {len(sections)} clusters"
    if generated_at:
        meta += f" · {generated_at}"
    lines.append(("meta", meta))
    for cluster, cell_type, confidence, verify, text in sections:
        head = f"{cluster} · {cell_type} — {confidence}" + ("  [re-check]" if verify else "")
        lines.append(("h2", head))
        lines.extend(_blocks(text))
    if global_check.strip():
        lines.append(("h2", "Cross-cluster global check"))
        lines.extend(_blocks(global_check))
    if caveats.strip():
        lines.append(("h2", "Caveats"))
        lines.extend(_blocks(caveats))
    if lab_notes:
        lines.append(("h2", "Lab notes (dataset / lab-wide)"))
        for nl in lab_notes:
            lines.append(("note", _note_text(nl)))
    return lines


def _short_set(gene_set: str) -> str:
    return gene_set.replace("HALLMARK_", "").replace("_", " ").title()


def global_check_text(holistic, themes) -> str:
    """Cross-cluster global check: does the whole set cohere? Marker coherence +
    the one refinement + the enrichment programs that recur across clusters. Every
    statement is read off the holistic review / themes (no new claim)."""
    paras: list[str] = []
    if holistic is not None:
        paras.extend(holistic.coherence_notes)
        for r in holistic.refinements:
            paras.append(
                f"Refinement to consider — {r.cluster}: {r.from_call} -> {r.to_call} ({r.rationale})."
            )
    if themes is not None and getattr(themes, "recurring", ()):
        bits = [
            f"{_short_set(t.gene_set)} (in {t.n_clusters} clusters: {', '.join(t.clusters)})"
            for t in themes.recurring[:4]
        ]
        paras.append(
            "Programs shared across clusters — " + "; ".join(bits) + ". A program enriched in "
            "many clusters partly reflects the panel's design and shared genes; weight it less "
            "as cluster-specific evidence."
        )
    return "\n\n".join(paras) if paras else "The set is coherent across clusters."


def caveats_text(verdicts, enrichments: dict, panel_size: int) -> str:
    """The interpretation's caveats: panel scope, panel-absence, flagged/small calls,
    and the cross-lineage-enrichment reminder. Derived deterministically from the
    verdicts + enrichment records; nothing invented."""
    paras = [
        f"Panel scope — this is a targeted {panel_size}-gene panel. Enrichment is measured only "
        "over each set's genes that are on the panel, so no program is a genome-wide claim, and a "
        "program can be missed simply because its genes were not on the panel.",
        "Panel absence — the absence of a canonical marker is not evidence against a cell type "
        "when that gene was never on the panel; missing canonical markers were checked against "
        "the panel before any down-weight.",
    ]
    flagged = [f"{v.cluster} {v.cell_type}" for v in verdicts if v.verify]
    if flagged:
        paras.append(
            "Flagged for re-check — " + ", ".join(flagged) + ". These calls rest on thin or "
            "non-specific evidence; confirm before relying on them."
        )
    small = [f"{v.cluster} {v.cell_type}" for v in verdicts if getattr(v, "small_n", False)]
    if small:
        paras.append(
            "Small clusters — " + ", ".join(small) + " have few assigned markers, so both the "
            "call and its enrichment are lower-powered."
        )
    paras.append(
        "Cross-lineage programs — a program enriched in a cluster whose leading-edge genes mark "
        "another lineage usually reflects co-infiltration or shared panel genes, not a re-typing "
        "of the cluster; these are flagged in the per-cluster programs above."
    )
    return "\n\n".join(paras)


def working_docx(sections, **kw) -> bytes:
    """Export the (edited) working space as a Word .docx."""
    return _lines_to_docx(_working_lines(sections, **kw))


def working_pdf(sections, **kw) -> bytes:
    """Export the (edited) working space as a PDF."""
    return _lines_to_pdf(_working_lines(sections, **kw))


__all__ = [
    "NoteLine",
    "ClusterSection",
    "ReportModel",
    "build_report",
    "build_report_from_sources",
    "report_html",
    "render_report_html",
    "report_to_docx",
    "report_to_pdf",
    "default_cluster_summary",
    "global_check_text",
    "caveats_text",
    "working_docx",
    "working_pdf",
]
